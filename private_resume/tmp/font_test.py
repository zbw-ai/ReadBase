from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


fonts = [
    "Noto Sans CJK SC",
    "STSong",
    "SimSong",
    "Songti SC",
    "STHeiti",
    "Heiti SC",
    "Hiragino Sans GB",
    "Arial Unicode MS",
    "PingFang SC",
]

doc = Document()
for name in fonts:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}: 曾柏炜 大模型训练推理工程师 中文测试 123 ABC")
    r.font.name = name
    r.font.size = Pt(14)
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    r._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)

doc.save("/Users/zengbw/ReadBase/private_resume/tmp/font-test.docx")
