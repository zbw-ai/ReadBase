from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


OUTPUT = Path(
    "/Users/zengbw/ReadBase/private_resume/output/"
    "曾柏炜-大模型训练推理Infra高级工程师-原格式版-2026.docx"
)
PHOTO = Path("/Users/zengbw/ReadBase/private_resume/tmp/original-photo.png")

FONT = "Arial Unicode MS"
BLACK = "000000"
BLUE = "4B83F5"
BADGE_FILL = "EAF2FF"
CONTENT_WIDTH_IN = 7.23


def set_run_font(run, size=8.2, bold=False, color=BLACK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rpr.rFonts.set(qn(key), FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "en-US")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size=8.2, bold=False):
    style.font.name = FONT
    rpr = style._element.get_or_add_rPr()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rpr.rFonts.set(qn(key), FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "en-US")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def make_picture_floating(run, x_in, y_in):
    inline = run._r.xpath("w:drawing/wp:inline")[0]
    anchor = OxmlElement("wp:anchor")
    for key, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(key, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    offset_h = OxmlElement("wp:posOffset")
    offset_h.text = str(int(Inches(x_in)))
    pos_h.append(offset_h)
    anchor.append(pos_h)

    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    offset_v = OxmlElement("wp:posOffset")
    offset_v.text = str(int(Inches(y_in)))
    pos_v.append(offset_v)
    anchor.append(pos_v)

    for tag in ("wp:extent", "wp:effectExtent"):
        element = inline.find(qn(tag))
        if element is not None:
            anchor.append(element)

    wrap = OxmlElement("wp:wrapNone")
    anchor.append(wrap)

    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        element = inline.find(qn(tag))
        if element is not None:
            if tag == "wp:docPr":
                element.set("descr", "曾柏炜证件照")
            anchor.append(element)

    inline.getparent().replace(inline, anchor)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Inches(0.52)
    section.right_margin = Inches(0.52)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.2)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, 8.2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.widow_control = True

    for style_name in ("Header", "Footer"):
        style = doc.styles[style_name]
        set_style_font(style, 8.0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

    # Keep the original template's first-page portrait in the header layer.
    hp = section.first_page_header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.line_spacing = 1.0
    picture_run = hp.add_run()
    picture_run.add_picture(str(PHOTO), width=Inches(0.91))
    make_picture_floating(picture_run, x_in=6.84, y_in=0.27)


def add_bottom_border(paragraph, size=7):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLACK)
    pbdr.append(bottom)


def add_header_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run("曾柏炜"), size=19.5, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run("☎  18859272673    ✉  zbw20@tsinghua.org.cn    ⌖  深圳"), size=9.3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run("▣  29岁    ♂  男    ⌂  福建 莆田    ◇  汉族"), size=9.1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7.0)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run("♙  大模型训练推理 Infra 高级工程师    ▣  SFT/RLVR/Agentic RL/长上下文/在线蒸馏"), size=8.9)


def add_section_heading(doc, text, before=4.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size=11.7, bold=True)
    add_bottom_border(p)
    return p


def set_right_tab(paragraph):
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )


def shade_run(run, fill=BADGE_FILL):
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rpr.append(shd)


def add_badge(paragraph, text):
    run = paragraph.add_run(f" {text} ")
    set_run_font(run, size=8.2, color=BLUE)
    shade_run(run)


def add_title_date(doc, title, date, size=9.2, after=0.3):
    p = doc.add_paragraph()
    set_right_tab(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(title), size=size, bold=True)
    set_run_font(p.add_run(f"\t{date}"), size=8.5)
    return p


def add_education(doc, school, detail, date, badges=None, honor=None):
    p = doc.add_paragraph()
    set_right_tab(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.2)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = honor is not None
    set_run_font(p.add_run(f"{school} {detail}"), size=9.1, bold=True)
    for badge in badges or []:
        add_badge(p, badge)
    set_run_font(p.add_run(f"\t{date}"), size=8.5)
    if honor:
        q = doc.add_paragraph()
        q.paragraph_format.space_before = Pt(0)
        q.paragraph_format.space_after = Pt(0.2)
        q.paragraph_format.line_spacing = 1.0
        set_run_font(q.add_run(honor), size=8.3)


def add_real_bullet_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "430")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "430")
    ind.set(qn("w:hanging"), "260")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "8")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), FONT)
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)


def add_bullet(doc, num_id, label, text, size=8.15):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.35)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    set_run_font(p.add_run(f"{label}："), size=size, bold=True)
    set_run_font(p.add_run(text), size=size)
    return p


def add_hyperlink(paragraph, text, url, size=8.15):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run._element.get_or_add_rPr().append(underline)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return run


def add_link_bullet(doc, num_id, label, link_text, url, size=8.15):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0.35)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    set_run_font(p.add_run(f"{label}："), size=size, bold=True)
    add_hyperlink(p, link_text, url, size=size)
    return p


def add_labeled_line(doc, label, text, size=8.15, after=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    set_run_font(p.add_run(f"{label}："), size=size)
    set_run_font(p.add_run(text), size=size)
    return p


def add_skill_bullet(doc, num_id, label, text):
    return add_bullet(doc, num_id, label, text, size=8.25)


def build():
    doc = Document()
    configure_document(doc)
    num_id = add_real_bullet_definition(doc)

    core = doc.core_properties
    core.title = "曾柏炜 - 大模型训练推理 Infra 高级工程师简历"
    core.subject = "按原简历格式更新的小鹏与华为工作经历"
    core.author = "曾柏炜"
    core.keywords = "大模型Infra, VeRL, AReaL, Megatron-Core, RLVR, Agentic RL"

    add_header_block(doc)

    add_section_heading(doc, "教育经历", before=0)
    add_education(
        doc,
        "清华大学",
        "电子信息 工学硕士 人工智能方向",
        "2020年09月 - 2023年06月",
        badges=["985", "211", "双一流"],
        honor="荣誉：2022年清华大学 校级一等奖学金（专业仅一人获得）",
    )
    add_education(
        doc,
        "厦门大学",
        "电气工程及其自动化 工学学士",
        "2015年09月 - 2019年06月",
        badges=["985", "211", "双一流"],
    )

    add_section_heading(doc, "工作技能", before=3.0)
    add_skill_bullet(
        doc,
        num_id,
        "模型训练",
        "了解Megatron 5D并行与显存优化机制，了解ZeRO优化原理与FSDP分片策略，熟悉AReaL跨引擎权重同步机制。",
    )
    add_skill_bullet(
        doc,
        num_id,
        "大模型后训练",
        "SFT、RLVR、Agentic RL、PPO/GRPO/DAPO、On-Policy Distillation/MOPD。",
    )
    add_skill_bullet(
        doc,
        num_id,
        "训练/推理框架",
        "PyTorch、Megatron-Core、VeRL、AReaL、vLLM、SGLang、Ray、MindSpeed、DeepSpeed。",
    )
    add_skill_bullet(
        doc,
        num_id,
        "性能优化",
        "fully-async policy、MoE、packed sequence、recompute/offload、CUDA Graph、XCCL/NCCL、MFU、显存建模、通信重叠、tracing、checkpoint/recovery。",
    )
    add_skill_bullet(
        doc,
        num_id,
        "工程能力",
        "Python、C++、Shell、Linux、Git；多机 GPU 与昇腾集群部署、性能分析及生产故障定位；英语 CET-6。",
    )

    add_section_heading(doc, "工作经历", before=3.2)
    add_title_date(
        doc,
        "小鹏机器人-大模型训练推理Infra高级工程师",
        "2025年11月 - 至今",
    )
    add_labeled_line(
        doc,
        "核心方向",
        "面向General/Math/Code等训练任务，负责基于VeRL、Megatron-Core和AReaL的训练推理Infra建设，主攻Qwen3/Qwen3.5 dense/MoE、32K-256K长上下文及多轮Agentic RL，覆盖SFT、RLVR与vLLM/SGLang rollout。",
    )
    add_labeled_line(
        doc,
        "工程交付",
        "交付Qwen3.5-35B-A3B 128K online RL基线、35B-MoE 256K及27B 128K/256K SFT checkpoint，沉淀多机GPU部署、MFU/吞吐看板与故障恢复能力。",
    )
    add_labeled_line(
        doc,
        "性能收益",
        "Qwen3.5-9B SFT单步由31s降至9.3s（3.3x），MFU由23%提升至29.6%；fully async RLVR代表性稳态吞吐由76提升至211-255 tokens/s/GPU。",
    )
    add_labeled_line(
        doc,
        "模型效果",
        "采用TMax-27B单Teacher时Terminal/SWE分别提升7.9pp/7.0pp；进一步验证双Teacher MOPD在SWE、Terminal取得双域提升，且General评测性能不下降。",
    )

    add_title_date(
        doc,
        "华为-计算产品线-昇腾计算-AI训练优化工程师",
        "2023年07月 - 2025年11月",
    )
    add_labeled_line(
        doc,
        "项目管理",
        "TX项目训练负责人，端到端负责客户训练解决方案设计和项目落地，包括TX项目TEG平台、PCG、CSIG、AI LAB、网平五个部门昇腾训练需求洞察、解决方案设计和保障客户训练业务长稳运行，带领团队4-5人。",
    )
    add_labeled_line(
        doc,
        "模型需求",
        "对齐客户需求，完成模型适配、精度对齐、性能优化到xx倍竞品，及现场POC演示产品功能性能。",
    )
    add_labeled_line(
        doc,
        "现网问题",
        "识别需求和风险，协调人力快速闭环问题，推动问题解决，以及赋能客户学会昇腾。",
    )
    add_labeled_line(
        doc,
        "模型经历",
        "HunyuanVideo、HunyuanDiT、星火MoE、星火图文系列、OpenSoraPlan等。",
    )
    add_labeled_line(
        doc,
        "工作成果",
        "xx局点xxB-MOE模型在x千卡集群长稳训练两个月、xx局点x万卡集群顺利交付和业务正常运行、xx推荐推理首单。",
    )

    add_title_date(doc, "微软亚洲研究院 - 机器学习组 - 算法实习生", "2022年07月 - 2022年10月")
    add_labeled_line(doc, "项目内容", "基于生成对抗网络实现金融时间序列（股票数据）的生成研究。")
    add_labeled_line(
        doc,
        "我的工作",
        "① 通过预测股票未来价格的方式，分别采用RNN作为生成器和判别器，在相互博弈中训练模型，模拟生成股票价格数据；② 通过分析自相关性、厚尾分布、波动率聚集、杠杆效应、粗细波动率相关、盈亏不对称性等评价指标，验证生成数据与真实数据具有相同统计特性；③ 为量化策略开发中提供更多训练样本和检验过拟合。",
    )

    project_heading = add_section_heading(doc, "项目经历", before=3.2)
    add_title_date(doc, "基于AReaL框架的Agentic RL与在线蒸馏", "2026年03月 - 2026年08月")
    add_bullet(
        doc,
        num_id,
        "长上下文与多轮Agentic RL",
        "面向Qwen3.5-9B 128K长上下文与多轮环境交互场景，搭建Agentic RL训练链路，打通多机GPU训练、vLLM/SGLang rollout及workflow/reward闭环。",
    )
    add_bullet(
        doc,
        num_id,
        "OPD/MOPD模型效果",
        "在9B模型上打通Math、SWE、Terminal三域；Math OPD在MATH500由69.80提升至75.89（+6.09pp）；采用TMax-27B单Teacher时Terminal/SWE分别提升7.9pp/7.0pp，进一步验证双Teacher MOPD实现SWE、Terminal双域提升且General评测不下降。",
    )

    add_title_date(doc, "基于VeRL框架的SFT/RLVR训练", "2025年12月 - 2026年03月")
    add_bullet(
        doc,
        num_id,
        "长上下文训练交付",
        "面向General/Math/Code等任务及Qwen3/Qwen3.5 dense/MoE，打通32K-256K长上下文SFT/RLVR训练链路，交付Qwen3.5-35B-A3B 128K online RL基线、35B-MoE 256K及27B 128K/256K SFT checkpoint，支持多机GPU训练与vLLM/SGLang rollout。",
    )
    add_bullet(
        doc,
        num_id,
        "Rollout与异步RLVR优化",
        "面向Qwen3-30B-A3B 32K训练场景重构gen-TP、实例数及Trainer/Rollouter资源配比，将代表性稳态吞吐由76提升至211-255 tokens/s/GPU；1:1配比达到236-293，idle ratio由0.41降至0.10-0.14；35B真实RL中CUDA Graph将decode提速约14x。",
    )
    add_bullet(
        doc,
        num_id,
        "SFT与长序列性能优化",
        "Qwen3.5-9B SFT通过num_workers 0->8与选择性重计算将step time由31s降至9.3s（3.3x），MFU由23%提升至29.6%；Qwen3.5-35B 128K平均耗时优化2x至2240s，并修复CP chunking静默失效导致7.6GB全量logits buffer分配问题。",
    )
    tx_project_title = add_title_date(doc, "TX项目训练业务", "2024年10月 - 2025年03月")
    tx_project_title.paragraph_format.page_break_before = True
    add_bullet(
        doc,
        num_id,
        "模型需求",
        "负责HunyuanVideo模型、参与HunyuanDiT、HunyuanLarge MOE模型的迁移适配、精度对齐、性能优化，通过优化并行策略、优化内存、使能融合算子、优化算子性能等手段，从开箱性能提升30-50%；分析HCCS和ROCE组网下8机性能及优化，进一步探索超大模型分布式训练并行策略研究。",
    )
    add_bullet(
        doc,
        num_id,
        "项目管理",
        "负责TEG平台、PCG、CSIG、AILAB、网平模型业务接口和整体看护，管理团队4-5人，保障模型需求交付、客户适配模型（5+），解决问题（80+）。",
    )

    add_title_date(doc, "X1项目训练业务", "2023年12月 - 2024年09月")
    add_bullet(
        doc,
        num_id,
        "MOE",
        "负责MOE-V2-xxxB模型功能打通，精度对齐，性能优化，从开箱0.16x优化到0.95x，MFU 35%达成客户目标后上线，并保障客户3K卡集群训练任务顺利完成。",
    )
    add_bullet(
        doc,
        num_id,
        "多模态",
        "负责图文类模型（图文ViT-13B/70B/xxxB、图文OCR、图文Swin）和类SORA模型（Vqvae、Stdit2、HunyuanDiT模型）功能打通、精度对齐、性能优化达标；看护模型训练场景，解决30+现网问题，保障千卡集群训练任务长稳进行。",
    )

    add_title_date(doc, "MT项目训练业务", "2023年07月 - 2023年11月")
    add_bullet(
        doc,
        num_id,
        "MT推荐模型推理业务精度调优",
        "通过分别Dump NPU和CPU数据进行Ait比对，对其算子精度，实现FP16模型精度达标、性能达标；达成客户预期，下单X卡；输出精度调优文档(x1)和工作周报(x5)。",
    )
    add_bullet(doc, num_id, "模型迁移", "YOLOv6训练loss跳变问题定位分析；LLAMA-13B模型迁移案例输出。")
    add_bullet(
        doc,
        num_id,
        "贵安AI硬装、软调实践环境解决方案",
        "训练、推理服务器配置方案，包括计算、存储、网络，设计服务器组网和设备清单配置。",
    )
    add_bullet(
        doc,
        num_id,
        "交付件",
        "撰写大模型训练组网方案，包括参数面网络、业务面/存储面网络、网络配置元组及参数面负载均衡方案，输出昇腾互联网行业大模型训练解决方案基线文档1份。",
    )

    add_section_heading(doc, "科研经历", before=4.0)
    add_title_date(
        doc,
        "基于半监督学习的乳腺组织PR免疫组化虚拟染色（项目负责人）",
        "2021年10月 - 2022年03月",
        size=8.9,
    )
    add_bullet(
        doc,
        num_id,
        "方法创新",
        "设计基于图像配准与相邻切片弱监督的乳腺病理虚拟染色框架，引入分类一致性约束，在无需专家像素级标注的条件下实现H&E至PR图像生成。",
    )
    add_bullet(
        doc,
        num_id,
        "论文成果",
        "第一作者论文发表于医学图像计算领域顶级会议MICCAI 2022（CCF-B），并获Student Travel Award。",
    )
    add_link_bullet(
        doc,
        num_id,
        "论文",
        "MICCAI 2022（Springer）",
        "https://link.springer.com/chapter/10.1007/978-3-031-16434-7_23",
    )

    add_title_date(
        doc,
        "基于无监督学习的肾组织特殊染色图像生成（核心成员）",
        "2020年12月 - 2021年09月",
        size=8.9,
    )
    add_bullet(
        doc,
        num_id,
        "方法创新",
        "提出多域虚拟染色GAN，通过风格编码动态建模染色域特征，并以正则化约束平衡组织结构保真与目标染色风格，首次实现多种肾组织特殊染色图像的相互转换。",
    )
    add_bullet(
        doc,
        num_id,
        "论文成果",
        "共同第一作者论文发表于人工智能领域顶级会议AAAI 2022（CCF-A）。",
    )
    add_link_bullet(
        doc,
        num_id,
        "论文",
        "AAAI 2022（AAAI Digital Library）",
        "https://ojs.aaai.org/index.php/AAAI/article/view/20054",
    )

    add_section_heading(doc, "竞赛经历", before=4.0)
    add_title_date(
        doc,
        "Kaggle：Classify Leaves -- Train models to predict the plant species（14th）",
        "2021年05月 - 2021年06月",
        size=8.9,
    )
    add_title_date(
        doc,
        "NLPCC Workshop 2021：Automatic Information Extraction（Top 3）",
        "2021年06月",
        size=8.9,
    )

    doc.settings.update_fields_on_open = True
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
