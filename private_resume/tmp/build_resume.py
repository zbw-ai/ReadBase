from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    "/Users/zengbw/ReadBase/private_resume/output/"
    "曾柏炜-大模型训练推理Infra高级工程师-2026.docx"
)


# Base preset: compact_reference_guide. Named resume overrides keep the document
# ATS-friendly and within a deliberate two-page budget.
TOKENS = {
    "page": {
        "width": 8.5,
        "height": 11.0,
        "margins": 1.0,
        "header_footer": 0.492,
        "content_width": 6.5,
    },
    "fonts": {"cjk": "Noto Sans CJK SC", "latin": "Noto Sans CJK SC"},
    "colors": {
        "navy": "17324D",
        "teal": "087E8B",
        "ink": "202A33",
        "muted": "59636E",
        "light": "D9E3E8",
    },
    "resume_override": {
        "body_size": 9.35,
        "body_line_spacing": 1.14,
        "body_after": 0.0,
        "section_size": 11.2,
        "section_before": 8.0,
        "section_after": 3.0,
        "bullet_marker": 0.18,
        "bullet_text": 0.38,
        "bullet_hanging": 0.20,
        "bullet_after": 2.6,
        "bullet_line_spacing": 1.12,
    },
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, color=None, bold=None, italic=None, font=None):
    font = font or TOKENS["fonts"]["cjk"]
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), TOKENS["fonts"]["latin"])
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), TOKENS["fonts"]["latin"])
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color, bold=False, italic=False):
    style.font.name = TOKENS["fonts"]["cjk"]
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), TOKENS["fonts"]["latin"])
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), TOKENS["fonts"]["latin"])
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), TOKENS["fonts"]["cjk"])
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style.font.italic = italic


def set_cell_free_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(TOKENS["page"]["width"])
    section.page_height = Inches(TOKENS["page"]["height"])
    section.top_margin = Inches(TOKENS["page"]["margins"])
    section.bottom_margin = Inches(TOKENS["page"]["margins"])
    section.left_margin = Inches(TOKENS["page"]["margins"])
    section.right_margin = Inches(TOKENS["page"]["margins"])
    section.header_distance = Inches(TOKENS["page"]["header_footer"])
    section.footer_distance = Inches(TOKENS["page"]["header_footer"])

    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(
        normal,
        TOKENS["resume_override"]["body_size"],
        TOKENS["colors"]["ink"],
    )
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(TOKENS["resume_override"]["body_after"])
    normal.paragraph_format.line_spacing = TOKENS["resume_override"]["body_line_spacing"]
    normal.paragraph_format.widow_control = True

    title = styles.add_style("Resume Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, 24, TOKENS["colors"]["navy"], bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1.5)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles.add_style("Resume Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle, 11.2, TOKENS["colors"]["teal"], bold=True)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(3.0)
    subtitle.paragraph_format.line_spacing = 1.0

    h1 = styles["Heading 1"]
    set_style_font(
        h1,
        TOKENS["resume_override"]["section_size"],
        TOKENS["colors"]["teal"],
        bold=True,
    )
    h1.paragraph_format.space_before = Pt(TOKENS["resume_override"]["section_before"])
    h1.paragraph_format.space_after = Pt(TOKENS["resume_override"]["section_after"])
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, 10.4, TOKENS["colors"]["navy"], bold=True)
    h2.paragraph_format.space_before = Pt(4.0)
    h2.paragraph_format.space_after = Pt(1.0)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3, 9.6, TOKENS["colors"]["navy"], bold=True)
    h3.paragraph_format.space_before = Pt(2.0)
    h3.paragraph_format.space_after = Pt(1.0)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    for style_name, size in (("Header", 8.0), ("Footer", 8.0)):
        style = styles[style_name]
        set_style_font(style, size, TOKENS["colors"]["muted"])
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0


def add_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "547")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "547")
    ind.set(qn("w:hanging"), "288")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "52")
    spacing.set(qn("w:line"), "269")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)

    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), TOKENS["fonts"]["latin"])
    rfonts.set(qn("w:hAnsi"), TOKENS["fonts"]["latin"])
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TOKENS["colors"]["teal"])
    rpr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    rpr.append(size)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8, color=TOKENS["colors"]["muted"])


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("曾柏炜  |  大模型训练推理 Infra  |  ")
    set_run_font(r, size=8, color=TOKENS["colors"]["muted"])
    add_field(p, "PAGE")
    r = p.add_run(" / ")
    set_run_font(r, size=8, color=TOKENS["colors"]["muted"])
    add_field(p, "NUMPAGES")


def add_contact_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5.0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("18859272673  |  zbw20@tsinghua.org.cn  |  深圳")
    set_run_font(r, size=9.0, color=TOKENS["colors"]["muted"])


def add_section(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True
    return p


def add_company_header(doc, company, org, date):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(TOKENS["page"]["content_width"]), WD_TAB_ALIGNMENT.RIGHT
    )
    r = p.add_run(company)
    set_run_font(r, size=10.4, color=TOKENS["colors"]["navy"], bold=True)
    if org:
        r = p.add_run(f"  |  {org}")
        set_run_font(r, size=9.35, color=TOKENS["colors"]["muted"], bold=False)
    r = p.add_run(f"\t{date}")
    set_run_font(r, size=9.15, color=TOKENS["colors"]["muted"], bold=True)
    return p


def add_role_line(doc, role, summary=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_with_next = True
    r = p.add_run(role)
    set_run_font(r, size=9.45, color=TOKENS["colors"]["teal"], bold=True)
    if summary:
        r = p.add_run(f"  |  {summary}")
        set_run_font(r, size=9.25, color=TOKENS["colors"]["ink"])
    return p


def add_bullet(doc, num_id, label, text):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(TOKENS["resume_override"]["bullet_after"])
    p.paragraph_format.line_spacing = TOKENS["resume_override"]["bullet_line_spacing"]
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    r = p.add_run(f"{label}：")
    set_run_font(r, size=9.35, color=TOKENS["colors"]["navy"], bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.35, color=TOKENS["colors"]["ink"])
    return p


def add_labeled_line(doc, label, text, after=1.7):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_together = True
    r = p.add_run(f"{label}：")
    set_run_font(r, size=9.25, color=TOKENS["colors"]["navy"], bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9.25, color=TOKENS["colors"]["ink"])
    return p


def add_education(doc, school, degree, date, detail=None):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(TOKENS["page"]["content_width"]), WD_TAB_ALIGNMENT.RIGHT
    )
    p.paragraph_format.space_before = Pt(1.0)
    p.paragraph_format.space_after = Pt(1.0)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    r = p.add_run(school)
    set_run_font(r, size=9.7, color=TOKENS["colors"]["navy"], bold=True)
    r = p.add_run(f"  |  {degree}")
    set_run_font(r, size=9.25, color=TOKENS["colors"]["ink"])
    r = p.add_run(f"\t{date}")
    set_run_font(r, size=9.05, color=TOKENS["colors"]["muted"], bold=True)
    if detail:
        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(0)
        d.paragraph_format.space_after = Pt(2.2)
        d.paragraph_format.line_spacing = 1.05
        r = d.add_run(detail)
        set_run_font(r, size=8.95, color=TOKENS["colors"]["muted"])


def build():
    doc = Document()
    set_cell_free_doc_defaults(doc)
    num_id = add_bullet_numbering(doc)
    add_footer(doc.sections[0])

    core = doc.core_properties
    core.title = "曾柏炜 - 大模型训练推理 Infra 高级工程师简历"
    core.subject = "大模型训练、推理、RLVR 与 Agentic RL Infra"
    core.author = "曾柏炜"
    core.keywords = "LLM Infra, verl, AReaL, Megatron-Core, RLVR, Agentic RL, SFT"
    core.comments = "面向大模型训练推理 Infra 高级工程师岗位的两页中文简历"

    p = doc.add_paragraph("曾柏炜", style="Resume Title")
    p.paragraph_format.keep_with_next = True
    p = doc.add_paragraph("大模型训练推理 Infra 高级工程师", style="Resume Subtitle")
    p.paragraph_format.keep_with_next = True
    add_contact_line(doc)

    add_section(doc, "职业概述")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3.0)
    p.paragraph_format.line_spacing = 1.14
    r = p.add_run("3 年大模型训练基础设施与性能优化经验")
    set_run_font(r, size=9.5, color=TOKENS["colors"]["navy"], bold=True)
    r = p.add_run(
        "，先后负责昇腾大规模训练优化与 GPU 后训练 Infra；当前聚焦基于 Megatron-Core、verl、AReaL 的长上下文 SFT/RLVR、异步 rollout、Agentic RL、轨迹数据链路和在线蒸馏。擅长从性能、数值正确性与模型效果三层建立可验证反馈回路。"
    )
    set_run_font(r, size=9.5, color=TOKENS["colors"]["ink"])

    add_section(doc, "核心能力")
    add_labeled_line(
        doc,
        "训练与 RL",
        "SFT、RLVR、PPO/GRPO、DAPO、On-Policy Distillation/MOPD、rule/model reward、policy staleness",
    )
    add_labeled_line(
        doc,
        "框架与推理",
        "PyTorch、Megatron-Core、verl、AReaL、vLLM、SGLang、Ray、MindSpeed、DeepSpeed",
    )
    add_labeled_line(
        doc,
        "分布式与性能",
        "TP/PP/CP/DP/EP、MoE、packed sequence、recompute/offload、NCCL/HCCS/RoCE、MFU、显存建模、通信重叠、tracing、checkpoint/recovery",
    )
    add_labeled_line(
        doc,
        "工程",
        "Python、C++、Shell、Linux、Git；多机 A100/昇腾集群部署、性能分析与生产故障定位",
        after=0.8,
    )

    add_section(doc, "工作经历")
    add_company_header(doc, "小鹏机器人", "大模型后训练基础设施", "2025.11 - 至今")
    add_role_line(
        doc,
        "大模型训练推理 Infra 高级工程师",
        "负责 SFT、RLVR 与 Agentic RL 框架建设及性能/正确性优化",
    )
    add_bullet(
        doc,
        num_id,
        "Fully Async RLVR",
        "面向 Qwen3-30B-A3B 32K、32 张 A100-80GB 场景适配 fully async 训练；基于 rollout 生产率与 Trainer 消费率重构 gen-TP、实例数及 Trainer/Rollouter 资源配比，将代表性稳态吞吐由 76 提升至 211-255 tokens/s/GPU；1:1 配比达到 236-293 tokens/s/GPU，Trainer idle ratio 由 0.41 降至 0.10-0.14。",
    )
    add_bullet(
        doc,
        num_id,
        "128K SFT",
        "打通 Qwen3.5 9B/27B 与 Qwen3 32B 在 16-64 张 A100 上的 128K 训练路径与边界验证；基于张量级显存账定位 fp32 logits、gradient buffer、CP 通信和 offload 开销，9B 场景采用 TP=2/CP=8，相比 TP=4/CP=4 将 step time 由约 163s 降至 102s（-37%），并补齐长样本 OOM、grad-norm spike 与 checkpoint deadlock 诊断。",
    )
    add_bullet(
        doc,
        num_id,
        "Agentic RL Rollout",
        "面向 Qwen3.5-9B 128K、32 cohorts x 8 trajectories、32 张 A100 场景建立 overlap-aware step、cohort tail、prefix cache 与 per-turn LLM 指标体系；量化基线 step 均值 83.89 min，其中 rollout wait 73.21 min、占 87.27%，并以 wait-after-7th p95 约 51.5 min 定位长上下文后期推理与 8-way cohort straggler 为一阶瓶颈。",
    )

    page_break = doc.add_paragraph()
    page_break.paragraph_format.space_before = Pt(0)
    page_break.paragraph_format.space_after = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    add_company_header(doc, "小鹏机器人（续）", "大模型后训练基础设施", "2025.11 - 至今")

    add_bullet(
        doc,
        num_id,
        "轨迹利用与算法正确性",
        "打通 generated -> manager -> workflow -> trainer -> loss -> policy gradient 六层 lineage；真实 6-step run 中实现 96 条训练轨迹 100% exact join，识别 94 条 gradient-active 与 2 条 compact-filtered 轨迹，后者消耗 159,330 full-sequence tokens（占 3.91%）但不产生梯度，并将 stale、partial、waiting/final-drain 与 terminal waste 分开归因。",
    )
    add_bullet(
        doc,
        num_id,
        "多 Teacher 在线蒸馏",
        "设计并实现 OPD/MOPD，支持 trajectory 按数据域路由至对应 Teacher 计算 logp，覆盖 score 校验、mopd_pg loss、混域训练、equal-trajectory weighting、session drain、断点续训与 held-out paired evaluation；跑通 9B 单 Teacher 30-step 及双 Teacher canary，建立 FUNCTIONAL/NUMERIC/EFFICACY 分层验收门禁。",
    )
    add_company_header(doc, "华为", "计算产品线・昇腾计算", "2023.07 - 2025.11")
    add_role_line(
        doc,
        "AI 训练优化工程师",
        "负责客户大模型迁移适配、精度对齐、性能优化及大规模集群交付",
    )
    add_bullet(
        doc,
        num_id,
        "项目与团队负责",
        "担任 TX 项目训练负责人，端到端负责训练解决方案设计、现场 POC 与长稳运行，覆盖 TEG、PCG、CSIG、AI LAB、网平等部门，带领 4-5 人团队交付 5+ 模型、闭环 80+ 现网问题。",
    )
    add_bullet(
        doc,
        num_id,
        "大模型性能优化",
        "负责 HunyuanVideo，参与 HunyuanDiT、HunyuanLarge MoE 等模型迁移与优化；通过并行策略、显存、融合算子与下发优化，将开箱性能提升 30%-50%，并完成 HCCS/RoCE 组网下 8 机性能分析。",
    )
    add_bullet(
        doc,
        num_id,
        "MoE 训练优化",
        "完成超大规模 MoE 模型功能打通、精度对齐与性能优化，将相对性能由 0.16x 提升至 0.95x、MFU 达 35%，支撑客户 3,000 卡集群训练任务顺利完成。",
    )
    add_bullet(
        doc,
        num_id,
        "多模态与稳定性",
        "适配图文 ViT、OCR、Swin 及 VQ-VAE、ST-DiT2、HunyuanDiT 等多模态/视频生成模型，完成精度与性能达标，解决 30+ 训练问题，保障千卡集群长稳运行。",
    )
    add_bullet(
        doc,
        num_id,
        "集群与网络方案",
        "参与千卡至万卡训练业务交付，输出计算、存储、参数面与业务面网络方案，覆盖 HCCS/RoCE 组网、参数面负载均衡和训练故障定位；完成 LLAMA-13B 迁移案例及 YOLOv6 loss 跳变分析。",
    )

    add_section(doc, "教育经历")
    add_education(
        doc,
        "清华大学",
        "工学硕士｜电子信息（人工智能方向）",
        "2020.09 - 2023.06",
        "2022 年清华大学校级一等奖学金（专业唯一获奖者）",
    )
    add_education(
        doc,
        "厦门大学",
        "工学学士｜电气工程及其自动化",
        "2015.09 - 2019.06",
    )

    add_section(doc, "科研与竞赛")
    add_bullet(
        doc,
        num_id,
        "MICCAI 2022",
        "第一作者，半监督乳腺组织 PR 免疫组化虚拟染色，获 Student Travel Award。",
    )
    add_bullet(
        doc,
        num_id,
        "AAAI 2022",
        "共同第一作者，无监督肾组织多域特殊染色图像生成。",
    )
    add_bullet(
        doc,
        num_id,
        "竞赛",
        "Kaggle Classify Leaves 第 14 名；NLPCC 2021 Automatic Information Extraction Top 3。",
    )

    # Ensure every section uses the same footer and print settings.
    for section in doc.sections:
        add_footer(section) if not section.footer.paragraphs[0].text else None

    doc.settings.update_fields_on_open = True
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
